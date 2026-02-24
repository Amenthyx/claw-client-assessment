"""
AI Assistant Solutions — Personal Consultation
Premium client intake questionnaire generator.
Generates a professional DOCX with Part A (Private) and Part B (Enterprise).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ---------------------------------------------------------------------------
# Design tokens
# ---------------------------------------------------------------------------
TEAL = RGBColor(0, 102, 153)
DARK_TEAL = RGBColor(0, 80, 120)
DARK_GRAY = RGBColor(51, 51, 51)
MED_GRAY = RGBColor(120, 120, 120)
LIGHT_LINE = RGBColor(180, 180, 180)
WHITE = RGBColor(255, 255, 255)

TABLE_HEADER_HEX = "006699"
ALT_ROW_HEX = "F2F2F2"
HIGHLIGHT_BOX_HEX = "E8F4F8"
FIELD_BG_HEX = "F7F7F7"

CHECKBOX = "\u2610"
OUTPUT = r"C:\Users\Software Engineering\Desktop\AI_Agent_Client_Needs_Assessment.docx"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def shade_cell(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def shade_row(row, hex_color):
    for c in row.cells:
        shade_cell(c, hex_color)


def set_cell_borders(cell, color="006699", size="4"):
    """Set thin borders on a single cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------
def heading(doc, text, level=1):
    """Add a teal-coloured heading."""
    sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12)}
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = TEAL
        run.font.size = sizes.get(level, Pt(12))
        run.font.name = "Calibri"
    return p


def body(doc, text, bold=False, italic=False, space_after=Pt(6), color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = color or DARK_GRAY
    run.bold = bold
    run.italic = italic
    return p


def body_multi(doc, segments, space_after=Pt(6)):
    """Add a paragraph with multiple styled runs.
    segments = [(text, bold, italic, color), ...]
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    for text, bld, ital, clr in segments:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = clr or DARK_GRAY
        run.bold = bld
        run.italic = ital
    return p


def checkbox(doc, text, indent_cm=0.5):
    p = doc.add_paragraph()
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(f"{CHECKBOX}  {text}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_GRAY
    return p


def answer_line(doc, label="", width=65):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if label:
        r = p.add_run(f"{label}: ")
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        r.bold = True
        r.font.color.rgb = DARK_GRAY
    r2 = p.add_run("_" * width)
    r2.font.size = Pt(11)
    r2.font.color.rgb = LIGHT_LINE
    return p


def open_field(doc, lines=4):
    """A light-gray box with blank lines for free-text answers."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, FIELD_BG_HEX)
    set_cell_borders(cell, color="CCCCCC", size="4")
    cell.text = ""
    for _ in range(lines):
        cp = cell.add_paragraph("")
        cp.paragraph_format.space_after = Pt(2)
    # set width
    for row in tbl.rows:
        for c in row.cells:
            c.width = Inches(6.5)
    return tbl


def rating_item(doc, text, indent_cm=0.5):
    p = doc.add_paragraph()
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"___  {text}")
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_GRAY
    return p


def page_break(doc):
    doc.add_page_break()


def spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


# ---------------------------------------------------------------------------
# Table helper
# ---------------------------------------------------------------------------
def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header
    hdr = tbl.rows[0]
    shade_row(hdr, TABLE_HEADER_HEX)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        r.font.color.rgb = WHITE

    # data rows
    for ri, rd in enumerate(rows):
        row = tbl.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(row, ALT_ROW_HEX)
        for ci, val in enumerate(rd):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = "Calibri"
            r.font.color.rgb = DARK_GRAY

    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    return tbl


def highlight_box(doc, title, body_text, bg_hex=HIGHLIGHT_BOX_HEX):
    """Single-cell table used as a highlighted callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, bg_hex)
    set_cell_borders(cell, color="006699", size="6")
    cell.text = ""

    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Calibri"
    r.font.color.rgb = TEAL
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(body_text)
    r2.font.size = Pt(10)
    r2.font.name = "Calibri"
    r2.font.color.rgb = DARK_GRAY

    for row in tbl.rows:
        for c in row.cells:
            c.width = Inches(6.5)

    return tbl


# ===================================================================
#  SECTION BUILDERS
# ===================================================================

def build_cover(doc):
    spacer(doc, 6)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI Assistant Solutions")
    r.font.size = Pt(28)
    r.font.color.rgb = TEAL
    r.bold = True
    r.font.name = "Calibri"

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Personal Consultation")
    r2.font.size = Pt(22)
    r2.font.color.rgb = TEAL
    r2.bold = True
    r2.font.name = "Calibri"

    spacer(doc, 1)

    # Tagline
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Tell us about your world. We'll build the perfect AI assistant for you.")
    r3.font.size = Pt(14)
    r3.font.color.rgb = DARK_GRAY
    r3.italic = True
    r3.font.name = "Calibri"

    spacer(doc, 2)

    # Company
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Amenthyx \u2014 AI Automation Experts")
    r4.font.size = Pt(16)
    r4.font.color.rgb = DARK_TEAL
    r4.bold = True
    r4.font.name = "Calibri"

    # Date
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("February 2026")
    r5.font.size = Pt(14)
    r5.font.color.rgb = DARK_GRAY
    r5.font.name = "Calibri"

    spacer(doc, 4)

    # Confidential note
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run(
        "Confidential \u2014 Your answers help us build your personalized solution"
    )
    r6.font.size = Pt(9)
    r6.font.color.rgb = MED_GRAY
    r6.italic = True
    r6.font.name = "Calibri"


def build_welcome(doc):
    heading(doc, "Welcome", 1)

    welcome_text = (
        "Thank you for your interest in working with us. We are genuinely excited to learn "
        "about your world and find ways to make your day easier.\n\n"
        "This questionnaire takes about 15 minutes to complete. There are no wrong answers "
        "\u2014 we simply want to understand how you spend your time, what tools you already "
        "use, and where an AI assistant could make the biggest difference for you.\n\n"
        "Once we receive your completed questionnaire, our team will analyze your answers "
        "and come back to you with a tailored proposal within 48 hours. The proposal will "
        "include a clear recommendation, a transparent price, and a timeline for getting "
        "your assistant up and running.\n\n"
        "Everything you share with us is strictly confidential and will only be used to "
        "design your solution.\n\n"
        "If any question does not apply to you, feel free to skip it. If you are unsure "
        "about something, just write a short note and we will clarify it together during "
        "our follow-up call.\n\n"
        "We look forward to building something great for you."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(welcome_text)
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_GRAY

    spacer(doc, 1)

    # Warm sign-off
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run("Warm regards,\nThe Amenthyx Team")
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    r2.font.color.rgb = TEAL
    r2.italic = True


# ===================================================================
#  PART A: PRIVATE CLIENT
# ===================================================================

def build_part_a_header(doc):
    # Decorative divider line using a single-row table
    divider = doc.add_table(rows=1, cols=1)
    divider.style = "Table Grid"
    divider.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = divider.rows[0].cells[0]
    shade_cell(cell, TABLE_HEADER_HEX)
    set_cell_borders(cell, color=TABLE_HEADER_HEX, size="2")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PART A")
    r.font.size = Pt(20)
    r.font.color.rgb = WHITE
    r.bold = True
    r.font.name = "Calibri"
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Your Personal AI Assistant \u2014 Tell Us About You")
    r2.font.size = Pt(14)
    r2.font.color.rgb = WHITE
    r2.font.name = "Calibri"
    for row in divider.rows:
        for c in row.cells:
            c.width = Inches(6.5)
    spacer(doc, 1)


def build_a1(doc):
    heading(doc, "A1. About You", 1)

    body(doc, "1. Full Name:", bold=True)
    answer_line(doc)

    body(doc, "2. What do you do for work?", bold=True)
    answer_line(doc)

    body(doc, "3. How would you describe your typical day? (Check all that apply)", bold=True)
    for item in [
        "I spend a lot of time on emails",
        "I manage appointments and meetings",
        "I research things online frequently",
        "I handle invoices, bills, or finances",
        "I manage social media accounts",
        "I write content (articles, posts, reports)",
        "I coordinate with other people (family, team, clients)",
        "I travel frequently and need things organized",
        "I manage a property or rental business",
    ]:
        checkbox(doc, item)
    checkbox(doc, "Other: ___________________________")

    spacer(doc, 1)
    body(doc, "4. What frustrates you most in your daily routine?", bold=True)
    open_field(doc, lines=4)

    spacer(doc, 1)
    body(doc, "5. How many hours per week do you spend on repetitive tasks you wish someone else could handle?", bold=True)
    for item in [
        "Less than 2 hours",
        "2\u20135 hours",
        "5\u201310 hours",
        "More than 10 hours",
    ]:
        checkbox(doc, item)


def build_a2(doc):
    heading(doc, "A2. Your Digital Life", 1)

    body(doc, "1. Which messaging apps do you use daily? (Check all that apply)", bold=True)
    for item in ["WhatsApp", "Telegram", "iMessage", "Signal", "Discord", "SMS"]:
        checkbox(doc, item)
    checkbox(doc, "Other: ___________________________")

    spacer(doc, 1)
    body(doc, "2. Which email provider(s) do you use?", bold=True)
    for item in ["Gmail", "Outlook / Hotmail", "Yahoo", "ProtonMail", "Work email"]:
        checkbox(doc, item)
    checkbox(doc, "Other: ___________________________")

    spacer(doc, 1)
    body(doc, "3. Which calendar do you use?", bold=True)
    for item in ["Google Calendar", "Apple Calendar", "Outlook Calendar", "None"]:
        checkbox(doc, item)
    checkbox(doc, "Other: ___________________________")

    spacer(doc, 1)
    body(doc, "4. Do you use any of these tools?", bold=True)
    for item in [
        "Google Drive / Docs",
        "Dropbox",
        "Notion",
        "Evernote",
        "Trello",
        "Todoist",
        "Spotify",
        "Smart home devices (Alexa, Google Home, Philips Hue)",
        "Accounting software (QuickBooks, FreshBooks, etc.)",
        "Social media management tools",
        "None of these",
    ]:
        checkbox(doc, item)
    checkbox(doc, "Other: ___________________________")

    spacer(doc, 1)
    body(doc, "5. Where do you prefer to interact with your AI assistant?", bold=True)
    for item in [
        "WhatsApp (just text me!)",
        "Telegram",
        "Email",
        "A web dashboard",
        "Voice commands",
        "I don't know yet",
    ]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "6. Where should your AI assistant run?", bold=True)
    body(doc, "(The assistant needs a device that stays on. This can be your own hardware or a cloud server.)")
    for item in [
        "On my own computer (desktop or laptop that stays on)",
        "On a home server or NAS I already own",
        "On a Raspberry Pi or small device I have",
        "On a cloud server (we can set this up for you)",
        "I don\u2019t have hardware \u2014 I\u2019d like you to handle this (Managed Service)",
        "I\u2019m not sure \u2014 let\u2019s discuss",
    ]:
        checkbox(doc, item)


def build_a3(doc):
    heading(doc, "A3. What Would You Love Your AI Assistant To Do?", 1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "Imagine you had a personal assistant available 24/7. "
        "What would you ask them to do?"
    )
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_GRAY
    r.italic = True

    body(doc, "Rate each item from 1 (not interested) to 5 (I need this!):", bold=True)

    spacer(doc, 1)
    body(doc, "Daily Life", bold=True, color=TEAL)
    for item in [
        "Read and summarize my emails every morning",
        "Manage my calendar \u2014 schedule, remind, reschedule",
        "Give me a daily briefing (weather, news, to-dos)",
        "Help me plan trips and travel itineraries",
        "Track my expenses and send me weekly summaries",
        "Remind me of important dates and follow-ups",
    ]:
        rating_item(doc, item)

    spacer(doc, 1)
    body(doc, "Communication", bold=True, color=TEAL)
    for item in [
        "Auto-reply to routine messages when I'm busy",
        "Draft professional emails based on my notes",
        "Translate messages in real-time",
        "Send scheduled messages to contacts",
        "Manage group chats or community channels",
    ]:
        rating_item(doc, item)

    spacer(doc, 1)
    body(doc, "Work & Productivity", bold=True, color=TEAL)
    for item in [
        "Research topics and give me summaries",
        "Write or edit documents, reports, or proposals",
        "Create presentations",
        "Monitor news or social media for specific topics",
        "Manage my files and organize documents",
    ]:
        rating_item(doc, item)

    spacer(doc, 1)
    body(doc, "Smart Home & Lifestyle", bold=True, color=TEAL)
    for item in [
        "Control my smart lights, thermostat, etc.",
        "Morning/evening routines automation",
        "Meal planning and recipe suggestions",
        "Fitness/health tracking reminders",
    ]:
        rating_item(doc, item)


# ---------------------------------------------------------------------------
# Shared 40-capability checklist
# ---------------------------------------------------------------------------
CAPABILITIES = [
    ("Your Emails & Messages", [
        "Reads your emails every morning and gives you a quick summary of what matters",
        "Drafts replies to routine emails so you just review and hit send",
        "Sorts your inbox automatically into urgent, newsletters, receipts, and spam",
        "Forwards the important stuff to your WhatsApp or Telegram right away",
        "Auto-replies to common questions when you\u2019re busy or on holiday",
        "Sends scheduled messages to contacts at exactly the right time",
    ]),
    ("Your Calendar & Schedule", [
        "Keeps your calendar organized \u2014 adds events, sends reminders, avoids double-bookings",
        "Finds free time slots and suggests meeting times to people for you",
        "Sends you a morning briefing with today\u2019s schedule, weather, and top priorities",
        "Reminds you of birthdays, deadlines, renewals, and follow-ups",
        "Syncs your personal and work calendars so nothing slips through",
    ]),
    ("Your Files & Documents", [
        "Keeps your Google Drive, Dropbox, or folders tidy and well-organized",
        "Reads contracts, PDFs, and long documents \u2014 gives you the key points in seconds",
        "Creates reports, summaries, and slide decks from your rough notes",
        "Backs up your important files automatically every day or week",
        "Converts documents between formats whenever you need (PDF, Word, Excel, etc.)",
    ]),
    ("Research & Staying Informed", [
        "Searches the web for you and delivers a clean, no-fluff summary",
        "Monitors news, competitors, or industry topics and sends you daily highlights",
        "Compares prices and options when you\u2019re shopping for products or services",
        "Tracks trends, mentions, or keywords across the web so you\u2019re always in the loop",
    ]),
    ("Social Media & Content", [
        "Writes ready-to-post social media captions, hashtags, and content ideas",
        "Schedules and publishes posts across all your social media accounts",
        "Watches your mentions and comments \u2014 alerts you when something needs attention",
        "Writes blog articles, newsletters, or marketing copy from just a few bullet points",
    ]),
    ("Money & Invoices", [
        "Tracks your spending from receipts, bank alerts, and invoices automatically",
        "Sends you a clear weekly or monthly budget summary",
        "Warns you before subscriptions renew so you can cancel what you don\u2019t need",
        "Creates professional invoices and sends payment reminders to clients",
    ]),
    ("Your Team & Customers", [
        "Answers common customer questions via chat or email around the clock",
        "Catches new leads and sends the best ones straight to you",
        "Sends polite follow-up emails after meetings so no opportunity gets forgotten",
        "Books appointments for clients and sends them automatic confirmations",
        "Summarizes meeting notes and distributes action items to your team",
        "Pulls together weekly performance reports from your business data",
    ]),
    ("Your Home & Daily Life", [
        "Controls your smart lights, thermostat, and appliances from a chat message",
        "Runs your morning and evening routines automatically (lights on, coffee, music, reminders)",
        "Sends you alerts from your security cameras or home sensors",
        "Plans your meals, suggests recipes, and builds your shopping list",
        "Organizes your trips \u2014 flights, hotels, things to do, all in one place",
        "Keeps you on track with fitness goals, health reminders, and habit streaks",
    ]),
]

CAPS_INTRO = (
    "Check everything that sounds useful to you. Don\u2019t worry about how it works "
    "\u2014 that\u2019s our job. Just tell us what you need."
)


def _build_capabilities_section(doc, section_heading):
    """Shared helper: emit a capabilities checklist with all 40 items."""
    heading(doc, section_heading, 1)
    body(doc, CAPS_INTRO, italic=True)
    for category, items in CAPABILITIES:
        spacer(doc, 1)
        heading(doc, category, 2)
        for item in items:
            checkbox(doc, item)


def build_a4_caps(doc):
    _build_capabilities_section(doc, "A4. Choose What Your AI Assistant Should Do")


def build_b3_caps(doc):
    _build_capabilities_section(doc, "B3. Choose What Your AI Assistant Should Do")


def build_a5_integration(doc):
    heading(doc, "A5. Integration & Automation", 1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "Which of your existing tools and services would you like your AI assistant "
        "to connect with and automate?"
    )
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_GRAY
    r.italic = True

    # Email & Communication
    spacer(doc, 1)
    body(doc, "Email & Communication", bold=True, color=TEAL)
    for item in [
        "Read and manage my Gmail / Outlook inbox",
        "Send emails on my behalf (with my approval)",
        "Auto-sort emails into categories (urgent, newsletters, receipts)",
        "Forward important emails to my WhatsApp / Telegram",
    ]:
        checkbox(doc, item)

    # Calendar & Scheduling
    spacer(doc, 1)
    body(doc, "Calendar & Scheduling", bold=True, color=TEAL)
    for item in [
        "Automatically add events from emails to my calendar",
        "Send me reminders before meetings",
        "Find free time slots and propose meetings",
        "Sync across multiple calendars",
    ]:
        checkbox(doc, item)

    # Files & Documents
    spacer(doc, 1)
    body(doc, "Files & Documents", bold=True, color=TEAL)
    for item in [
        "Organize files in my Google Drive / Dropbox",
        "Convert documents between formats",
        "Extract key information from PDFs and documents",
        "Backup important files automatically",
    ]:
        checkbox(doc, item)

    # Finance & Shopping
    spacer(doc, 1)
    body(doc, "Finance & Shopping", bold=True, color=TEAL)
    for item in [
        "Track my subscriptions and alert me before renewals",
        "Categorize my expenses from receipts / bank notifications",
        "Compare prices when I want to buy something",
        "Send me budget summaries",
    ]:
        checkbox(doc, item)

    # Social Media
    spacer(doc, 1)
    body(doc, "Social Media", bold=True, color=TEAL)
    for item in [
        "Post to my social media accounts on schedule",
        "Monitor mentions and comments",
        "Generate content ideas based on trending topics",
        "Track my followers and engagement",
    ]:
        checkbox(doc, item)

    # Smart Home
    spacer(doc, 1)
    body(doc, "Smart Home", bold=True, color=TEAL)
    for item in [
        "Control lights, heating, and appliances",
        "Set up morning / evening automation routines",
        "Security alerts from cameras / sensors",
        "Voice-activated commands via messaging app",
    ]:
        checkbox(doc, item)

    # Custom
    spacer(doc, 1)
    body(doc, "Custom Automations", bold=True, color=TEAL)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(
        "Is there something specific you do repeatedly that you'd love to automate? "
        "Describe it in your own words:"
    )
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    r2.font.color.rgb = DARK_GRAY
    r2.italic = True
    open_field(doc, lines=6)


def build_a6_privacy(doc):
    heading(doc, "A6. Privacy & Preferences", 1)

    body(doc, "1. How comfortable are you with your AI assistant accessing your data?", bold=True)
    for item in [
        "Full access \u2014 I want it to help with everything",
        "Moderate \u2014 It can read my calendar and emails, but not financial data",
        "Limited \u2014 Only what I explicitly share with it",
        "Minimal \u2014 I'll give it tasks manually each time",
    ]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "2. Should the assistant be available 24/7 or only during certain hours?", bold=True)
    checkbox(doc, "Always on")
    checkbox(doc, "Only during work hours")
    checkbox(doc, "Custom schedule: ___________________________")

    spacer(doc, 1)
    body(doc, "3. Will anyone else use this assistant besides you?", bold=True)
    checkbox(doc, "Just me")
    checkbox(doc, "My partner / family (how many? ___)")
    checkbox(doc, "My small team (how many? ___)")

    spacer(doc, 1)
    body(doc, "4. Any specific personality you'd like your assistant to have?", bold=True)
    body(doc, "(e.g., formal, casual, funny, minimalist, warm, direct)", italic=True, color=MED_GRAY)
    answer_line(doc)


# ===================================================================
#  PART B: ENTERPRISE CLIENT
# ===================================================================

def build_part_b_header(doc):
    divider = doc.add_table(rows=1, cols=1)
    divider.style = "Table Grid"
    divider.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = divider.rows[0].cells[0]
    shade_cell(cell, TABLE_HEADER_HEX)
    set_cell_borders(cell, color=TABLE_HEADER_HEX, size="2")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PART B")
    r.font.size = Pt(20)
    r.font.color.rgb = WHITE
    r.bold = True
    r.font.name = "Calibri"
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Your Business AI Assistant \u2014 Tell Us About Your Company")
    r2.font.size = Pt(14)
    r2.font.color.rgb = WHITE
    r2.font.name = "Calibri"
    for row in divider.rows:
        for c in row.cells:
            c.width = Inches(6.5)
    spacer(doc, 1)


def build_b1(doc):
    heading(doc, "B1. Company Profile", 1)

    body(doc, "1. Company Name:", bold=True)
    answer_line(doc)

    body(doc, "2. Your Name & Role:", bold=True)
    answer_line(doc)

    body(doc, "3. Industry:", bold=True)
    for item in [
        "Real Estate", "E-commerce / Retail", "Healthcare", "Finance / Banking",
        "Legal", "Marketing / Creative", "Technology / SaaS", "Education",
        "Hospitality / Tourism", "Manufacturing", "Consulting", "Logistics",
    ]:
        checkbox(doc, item)
    checkbox(doc, "Other: ___________________________")

    spacer(doc, 1)
    body(doc, "4. Number of Employees:", bold=True)
    for item in ["2\u201310", "11\u201350", "51\u2013200", "200\u20131,000", "1,000+"]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "5. How many departments would use the AI assistant?", bold=True)
    for item in [
        "Just mine",
        "2\u20133 departments",
        "Company-wide",
        "Not sure yet",
    ]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "6. Annual revenue range (helps us size the solution):", bold=True)
    for item in [
        "Under \u20ac100K",
        "\u20ac100K\u2013500K",
        "\u20ac500K\u20132M",
        "\u20ac2M\u201310M",
        "\u20ac10M+",
        "Prefer not to say",
    ]:
        checkbox(doc, item)


def build_b2(doc):
    heading(doc, "B2. Current Pain Points", 1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "What wastes the most time in your organization? "
        "Rate each from 1 (minor issue) to 5 (major bottleneck):"
    )
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_GRAY
    r.italic = True

    for item in [
        "Answering repetitive customer questions",
        "Manual data entry and report creation",
        "Scheduling and coordination between teams",
        "Email overload and slow response times",
        "Lead follow-up falling through the cracks",
        "Document review and approval processes",
        "Onboarding new employees",
        "Invoice processing and expense management",
        "Social media and marketing content",
        "IT support and troubleshooting",
        "Compliance and regulatory tasks",
        "Inventory and supply chain tracking",
    ]:
        rating_item(doc, item)

    spacer(doc, 1)
    body(doc, "Other pain points you'd like to mention:", bold=True)
    open_field(doc, lines=4)


def build_b4(doc):
    heading(doc, "B4. Integration & Automation Priorities", 1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "Which workflows would you like the AI assistant to automate? "
        "Rate each from 1 (low priority) to 5 (high priority):"
    )
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_GRAY
    r.italic = True

    # Customer-Facing
    body(doc, "Customer-Facing", bold=True, color=TEAL)
    for item in [
        "Answer customer questions via chat / email automatically",
        "Qualify leads and route to the right sales rep",
        "Send follow-up emails after meetings or inquiries",
        "Handle appointment booking for clients",
        "Process returns, refunds, or complaint tickets",
        "Collect customer feedback automatically",
    ]:
        rating_item(doc, item)

    # Internal Operations
    spacer(doc, 1)
    body(doc, "Internal Operations", bold=True, color=TEAL)
    for item in [
        "Generate weekly / monthly reports from your data",
        "Summarize meeting notes and distribute action items",
        "Automate invoice creation and send payment reminders",
        "Route internal requests to the right department",
        "Monitor key performance indicators and alert when something is off",
        "Automate employee onboarding checklists",
    ]:
        rating_item(doc, item)

    # Marketing & Sales
    spacer(doc, 1)
    body(doc, "Marketing & Sales", bold=True, color=TEAL)
    for item in [
        "Create and schedule social media posts",
        "Write email newsletters and campaigns",
        "Track campaign performance and generate reports",
        "Monitor competitor activity and industry news",
        "Generate product descriptions and marketing copy",
    ]:
        rating_item(doc, item)

    # Data & Documents
    spacer(doc, 1)
    body(doc, "Data & Documents", bold=True, color=TEAL)
    for item in [
        "Extract data from documents (invoices, contracts, forms)",
        "Keep databases and spreadsheets synchronized",
        "Generate formatted reports from raw data",
        "Ensure compliance documents are up to date",
        "Archive and organize company documents",
    ]:
        rating_item(doc, item)

    # IT & Development
    spacer(doc, 1)
    body(doc, "IT & Development (if applicable)", bold=True, color=TEAL)
    for item in [
        "Monitor servers and alert on issues",
        "Automate deployment and testing pipelines",
        "Manage code reviews and pull requests",
        "Track bugs and prioritize them",
    ]:
        rating_item(doc, item)

    # Custom
    spacer(doc, 1)
    body(doc, "Custom Workflows", bold=True, color=TEAL)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(
        "Describe any specific process unique to your business that you'd love to automate:"
    )
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    r2.font.color.rgb = DARK_GRAY
    r2.italic = True
    open_field(doc, lines=8)


def build_b5(doc):
    heading(doc, "B5. Compliance & Security", 1)

    body(doc, "1. What type of data will the AI assistant handle?", bold=True)
    for item in [
        "General business data (not sensitive)",
        "Customer personal data (names, emails, phones)",
        "Financial / payment data",
        "Health / medical records",
        "Legal / confidential documents",
        "Trade secrets / intellectual property",
    ]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "2. Compliance requirements:", bold=True)
    for item in ["GDPR", "HIPAA", "SOC 2", "PCI-DSS", "ISO 27001", "None / Not sure"]:
        checkbox(doc, item)
    checkbox(doc, "Industry-specific: ___________________________")

    spacer(doc, 1)
    body(doc, "3. Data hosting preference:", bold=True)
    for item in [
        "Must stay on our own servers (on-premise)",
        "Private cloud in EU",
        "Private cloud (any region)",
        "No preference",
    ]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "4. Available infrastructure:", bold=True)
    body(doc, "(The AI assistant needs hardware to run on. Do you already have something available?)")
    for item in [
        "We have our own servers (on-premise or data center)",
        "We already use cloud infrastructure (AWS, Azure, Google Cloud, etc.)",
        "We have a dedicated machine or NAS we can use",
        "We don\u2019t have infrastructure \u2014 we\u2019d like you to handle hosting (Managed Service)",
        "Not sure \u2014 let\u2019s discuss during the proposal",
    ]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "5. Who should approve AI actions before they are executed?", bold=True)
    for item in [
        "Nobody \u2014 fully autonomous is fine",
        "Manager approval for external actions (emails, messages to clients)",
        "Approval for all actions",
        "Depends on the action (we'll define rules together)",
    ]:
        checkbox(doc, item)


def build_b6(doc):
    heading(doc, "B6. Scale & Growth", 1)

    body(doc, "1. How many people will interact with the AI assistant daily?", bold=True)
    for item in ["1\u20135", "5\u201320", "20\u2013100", "100+"]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "2. Expected daily tasks for the AI assistant:", bold=True)
    for item in ["Less than 20", "20\u2013100", "100\u2013500", "500+"]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "3. How fast does it need to respond?", bold=True)
    for item in [
        "Instant (under 5 seconds)",
        "Quick (under 30 seconds)",
        "Background processing is fine",
    ]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "4. Growth plans in the next 12 months?", bold=True)
    for item in [
        "Stay the same",
        "Double our usage",
        "5x growth",
        "Planning rapid expansion",
    ]:
        checkbox(doc, item)


# ===================================================================
#  SECTION C: PRICING
# ===================================================================

def build_section_c(doc):
    # Section header
    divider = doc.add_table(rows=1, cols=1)
    divider.style = "Table Grid"
    divider.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = divider.rows[0].cells[0]
    shade_cell(cell, TABLE_HEADER_HEX)
    set_cell_borders(cell, color=TABLE_HEADER_HEX, size="2")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SECTION C")
    r.font.size = Pt(20)
    r.font.color.rgb = WHITE
    r.bold = True
    r.font.name = "Calibri"
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Service Packages & Pricing")
    r2.font.size = Pt(14)
    r2.font.color.rgb = WHITE
    r2.font.name = "Calibri"
    for row in divider.rows:
        for c in row.cells:
            c.width = Inches(6.5)
    spacer(doc, 1)

    # --- Private Clients ---
    heading(doc, "For Private Clients", 1)

    add_table(doc,
              ["", "Private Solution"],
              [
                  ["Investment", "\u20ac1,000 (one-time)"],
                  ["What's included", "Full AI assistant setup, configuration, and personalization"],
                  ["Hosting", "Runs on your own hardware (PC, server, Raspberry Pi) \u2014 or we set up cloud hosting for you"],
                  ["Channels", "All your messaging apps + email"],
                  ["Automations", "Custom workflows tailored to your needs"],
                  ["Scheduling", "Unlimited scheduled tasks and briefings"],
                  ["Smart Home", "Included if requested"],
                  ["Personality", "Fully customized to your preferences"],
                  ["Support", "Email support included during setup"],
                  ["Delivery", "48\u201372 hours"],
              ],
              col_widths=[2.2, 4.4])

    spacer(doc, 1)
    body(doc,
         "Note: The AI assistant requires an API subscription to an AI provider "
         "(e.g., Anthropic, OpenAI, or others). This is a separate cost managed "
         "directly by you, typically \u20ac5\u2013\u20ac50/month depending on usage. "
         "We will guide you through the setup.")

    spacer(doc, 2)

    # --- Enterprise Clients ---
    heading(doc, "For Enterprise Clients", 1)

    add_table(doc,
              ["", "Enterprise Solution"],
              [
                  ["Investment", "From \u20ac5,000 (one-time)"],
                  ["What's included", "Full deployment, integrations, custom workflows, team onboarding"],
                  ["Hosting", "Your own servers, your cloud, or we provide infrastructure"],
                  ["Users", "Unlimited"],
                  ["Channels", "All channels (chat, email, internal tools)"],
                  ["Integrations", "All your existing tools connected"],
                  ["Automations", "Unlimited custom workflows"],
                  ["Reporting", "Dashboards and automated reports"],
                  ["Compliance", "GDPR, SOC 2, industry-specific as needed"],
                  ["Support", "Dedicated account manager during setup"],
                  ["Delivery", "1\u20132 weeks depending on complexity"],
              ],
              col_widths=[2.2, 4.4])

    spacer(doc, 1)
    body(doc,
         "Note: Enterprise pricing starts at \u20ac5,000 and varies based on the number "
         "of integrations, custom workflows, and compliance requirements. "
         "API subscription costs are managed directly by your organization.")

    spacer(doc, 2)

    # --- Managed Service ---
    heading(doc, "Managed Service", 1)

    add_table(doc,
              ["", "Managed"],
              [
                  ["Price", "\u20ac300/month"],
                  ["Installation", "Included (no separate setup fee)"],
                  ["Hosting", "We provide and manage all infrastructure \u2014 or we manage it on your hardware"],
                  ["Updates & optimization", "Continuous, automatic"],
                  ["Monitoring", "24/7 health monitoring"],
                  ["Support", "Priority email and chat"],
                  ["Ideal for", "Clients who want zero hassle \u2014 whether on our servers or yours"],
              ],
              col_widths=[2.2, 4.4])

    spacer(doc, 2)

    # --- Ongoing Assistance ---
    heading(doc, "Ongoing Assistance (after 6 months)", 1)

    add_table(doc,
              ["", "Assistance"],
              [
                  ["Price", "\u20ac500/month"],
                  ["Available", "After the first 6 months of operation"],
                  ["Priority support", "Dedicated response within hours"],
                  ["Monthly optimization", "Performance review and improvement call"],
                  ["New integrations", "Connect new tools and services on request"],
                  ["Workflow updates", "Adapt automations as your needs evolve"],
                  ["Ideal for", "Growing businesses that need continuous evolution"],
              ],
              col_widths=[2.2, 4.4])

    spacer(doc, 2)

    # --- Understanding the Costs ---
    heading(doc, "Understanding the Costs", 1)

    body(doc,
         "Your AI assistant has two types of costs: our service fee (setup and management) "
         "and the AI provider subscription (like a phone plan for your assistant). "
         "Here\u2019s how it works:")

    body(doc,
         "Our fee covers everything we do: designing your assistant, configuring it, "
         "connecting your tools, and making sure it works perfectly. The AI provider fee "
         "is what you pay for the \u2018brain\u2019 of your assistant \u2014 this goes directly to "
         "companies like Anthropic or OpenAI, and depends on how much you use it.")

    spacer(doc, 1)

    # --- Cost Estimation Table ---
    heading(doc, "Estimated Monthly AI Provider Cost (Based on Your Usage)", 2)

    add_table(doc,
              ["Your Daily Usage", "Estimated Cost/Month", "What That Looks Like"],
              [
                  ["Light (5\u201310 tasks/day)", "\u20ac5\u2013\u20ac15",
                   "A few emails, calendar checks, daily briefing"],
                  ["Moderate (20\u201350 tasks/day)", "\u20ac15\u2013\u20ac40",
                   "Email management, scheduling, research, content drafts"],
                  ["Heavy (50\u2013100 tasks/day)", "\u20ac40\u2013\u20ac80",
                   "Full inbox management, team automation, reports"],
                  ["Intensive (100+ tasks/day)", "\u20ac80\u2013\u20ac200",
                   "Enterprise: customer support, lead qualification, multi-department"],
              ],
              col_widths=[2.0, 1.5, 3.0])

    spacer(doc, 1)
    body(doc,
         "These are approximate costs paid directly to the AI provider. "
         "We\u2019ll help you pick the most cost-effective option for your needs.",
         italic=True, color=MED_GRAY)

    spacer(doc, 1)

    # --- Total Cost Example 1: Private ---
    highlight_box(
        doc,
        "Example: Private client, moderate usage",
        "\u2022 Setup: \u20ac1,000 (one-time)\n"
        "\u2022 AI provider: ~\u20ac25/month\n"
        "\u2022 First year total: \u20ac1,000 + (\u20ac25 \u00d7 12) = \u20ac1,300\n"
        "\u2022 That\u2019s about \u20ac108/month for a 24/7 personal assistant",
        bg_hex=HIGHLIGHT_BOX_HEX,
    )

    spacer(doc, 1)

    # --- Total Cost Example 2: Enterprise ---
    highlight_box(
        doc,
        "Example: Enterprise with Managed Service",
        "\u2022 Managed service: \u20ac300/month (installation included)\n"
        "\u2022 AI provider: ~\u20ac60/month\n"
        "\u2022 Total: \u20ac360/month\n"
        "\u2022 For a team of 20, that\u2019s just \u20ac18 per person per month",
        bg_hex=HIGHLIGHT_BOX_HEX,
    )

    spacer(doc, 2)

    # --- Value Proposition Box ---
    highlight_box(
        doc,
        "Why This Pays for Itself",
        "Our clients typically save 10\u201320 hours per week on repetitive tasks. "
        "At an average rate of \u20ac50/hour, that's \u20ac2,000\u2013\u20ac4,000/month in recovered "
        "productivity \u2014 far exceeding the cost of the service.",
        bg_hex=HIGHLIGHT_BOX_HEX,
    )


# ===================================================================
#  SECTION D: AUTHORIZATION & NEXT STEPS
# ===================================================================

def build_section_d(doc):
    # Section header
    divider = doc.add_table(rows=1, cols=1)
    divider.style = "Table Grid"
    divider.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = divider.rows[0].cells[0]
    shade_cell(cell, TABLE_HEADER_HEX)
    set_cell_borders(cell, color=TABLE_HEADER_HEX, size="2")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SECTION D")
    r.font.size = Pt(20)
    r.font.color.rgb = WHITE
    r.bold = True
    r.font.name = "Calibri"
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Authorization & Next Steps")
    r2.font.size = Pt(14)
    r2.font.color.rgb = WHITE
    r2.font.name = "Calibri"
    for row in divider.rows:
        for c in row.cells:
            c.width = Inches(6.5)
    spacer(doc, 1)

    heading(doc, "Your Choice", 1)

    body(doc, "1. Which solution interests you?", bold=True)
    for item in [
        "Private (\u20ac1,000)",
        "Enterprise (from \u20ac5,000)",
        "Managed Service (\u20ac300/month \u2014 installation included)",
        "Not sure yet \u2014 let\u2019s discuss",
    ]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "2. Are you interested in Ongoing Assistance (\u20ac500/month, available after 6 months)?", bold=True)
    for item in ["Yes", "No", "Tell me more"]:
        checkbox(doc, item)

    spacer(doc, 1)
    body(doc, "3. Preferred start date:", bold=True)
    answer_line(doc)

    spacer(doc, 1)
    body(doc, "4. Anything else you'd like us to know?", bold=True)
    open_field(doc, lines=6)

    spacer(doc, 1)
    body(doc, "5. How did you hear about us?", bold=True)
    for item in [
        "Word of mouth",
        "Social media",
        "Google search",
        "LinkedIn",
        "Event or conference",
    ]:
        checkbox(doc, item)
    checkbox(doc, "Other: ___________________________")

    spacer(doc, 1)
    body(doc, "6. Authorization", bold=True)
    body(doc,
         "By signing below, you authorize our team to use the information provided in "
         "this questionnaire to design and build a tailored AI assistant solution on your behalf.")

    spacer(doc, 1)
    answer_line(doc, "Signature")
    answer_line(doc, "Date")

    spacer(doc, 2)

    # --- Contact info ---
    heading(doc, "Contact Information", 2)

    contact_tbl = doc.add_table(rows=1, cols=1)
    contact_tbl.style = "Table Grid"
    contact_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = contact_tbl.rows[0].cells[0]
    shade_cell(cell, HIGHLIGHT_BOX_HEX)
    set_cell_borders(cell, color="006699", size="6")
    cell.text = ""

    lines = [
        ("Amenthyx \u2014 AI Automation Experts", True, Pt(13), TEAL),
        ("", False, Pt(6), DARK_GRAY),
        ("GitHub: https://github.com/Amenthyx", False, Pt(10), DARK_GRAY),
        ("Deployment Toolkit: https://github.com/Amenthyx/claw-one-click-deploy", False, Pt(10), DARK_GRAY),
        ("Assessment Toolkit: https://github.com/Amenthyx/claw-client-assessment", False, Pt(10), DARK_GRAY),
    ]

    for i, (text, bld, size, color) in enumerate(lines):
        if i == 0:
            cp = cell.paragraphs[0]
        else:
            cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(text)
        cr.bold = bld
        cr.font.size = size
        cr.font.name = "Calibri"
        cr.font.color.rgb = color

    for row in contact_tbl.rows:
        for c in row.cells:
            c.width = Inches(6.5)


# ===================================================================
#  MAIN
# ===================================================================

def main():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GRAY

    for hs in ["Heading 1", "Heading 2", "Heading 3"]:
        if hs in doc.styles:
            doc.styles[hs].font.name = "Calibri"

    # Margins: 2cm all sides
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # === COVER ===
    build_cover(doc)
    page_break(doc)

    # === WELCOME ===
    build_welcome(doc)
    page_break(doc)

    # === PART A: PRIVATE CLIENT ===
    build_part_a_header(doc)
    build_a1(doc)
    page_break(doc)
    build_a2(doc)
    page_break(doc)
    build_a3(doc)
    page_break(doc)
    build_a4_caps(doc)
    page_break(doc)
    build_a5_integration(doc)
    page_break(doc)
    build_a6_privacy(doc)
    page_break(doc)

    # === PART B: ENTERPRISE CLIENT ===
    build_part_b_header(doc)
    build_b1(doc)
    page_break(doc)
    build_b2(doc)
    page_break(doc)
    build_b3_caps(doc)
    page_break(doc)
    build_b4(doc)
    page_break(doc)
    build_b5(doc)
    page_break(doc)
    build_b6(doc)
    page_break(doc)

    # === SECTION C: PRICING ===
    build_section_c(doc)
    page_break(doc)

    # === SECTION D: AUTHORIZATION ===
    build_section_d(doc)

    # Save
    doc.save(OUTPUT)
    size = os.path.getsize(OUTPUT)
    print(f"Document saved to: {OUTPUT}")
    print(f"File size: {size:,} bytes")


if __name__ == "__main__":
    main()
